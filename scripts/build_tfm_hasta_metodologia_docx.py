"""Build the TFM proposal through methodology as a DOCX deliverable."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "doc-01-proposal"
OUT_DOCX = OUT_DIR / "TFM_MayorgaTaborda_hasta_metodologia.docx"
LATEX_SOURCE = OUT_DIR / "main_hasta_metodologia.tex"
BIB_SOURCE = OUT_DIR / "references.bib"

TITLE = (
    "Control distribuido basado en dinamicas poblacionales para la formacion "
    "adaptativa de coaliciones en transporte cooperativo multi-AGV"
)
SUBTITLE = (
    "Propuesta de Trabajo Fin de Master hasta metodologia, con orientacion a "
    "continuidad doctoral y publicacion"
)
AUTHOR = "Jorge Luis Mayorga Taborda"
TUTOR = "Jose Ignacio Iniguez Amigot"
PROGRAM = "Master Universitario en Robotica y Automatizacion de Procesos"
INSTITUTION = "Universidad Internacional de Valencia (VIU)"
DATE_TEXT = "27 de mayo de 2026"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "F4F6F9"


def fix_text(text: str) -> str:
    """Apply Spanish accents after drafting with ASCII-safe source text."""

    replacements = {
        "MASTER": "MÁSTER",
        "Master": "Máster",
        "master": "máster",
        "Robotica": "Robótica",
        "robotica": "robótica",
        "Automatizacion": "Automatización",
        "automatizacion": "automatización",
        "dinamicas": "dinámicas",
        "dinamica": "dinámica",
        "formacion": "formación",
        "coaliciones": "coaliciones",
        "coalicion": "coalición",
        "metodologia": "metodología",
        "Metodologia": "Metodología",
        "orientacion": "orientación",
        "publicacion": "publicación",
        "Institucion": "Institución",
        "institucion": "institución",
        "Edicion": "Edición",
        "Iniguez": "Iñíguez",
        "Jose": "José",
        "Introduccion": "Introducción",
        "introduccion": "introducción",
        "justificacion": "justificación",
        "investigacion": "investigación",
        "Investigacion": "Investigación",
        "hipotesis": "hipótesis",
        "Hipotesis": "Hipótesis",
        "Objetivos especificos": "Objetivos específicos",
        "especificos": "específicos",
        "bibliograficas": "bibliográficas",
        "moviles": "móviles",
        "tamano": "tamaño",
        "coordinacion": "coordinación",
        "comunicacion": "comunicación",
        "informacion": "información",
        "linea": "línea",
        "lineas": "líneas",
        "validacion": "validación",
        "simulacion": "simulación",
        "comparacion": "comparación",
        "contribucion": "contribución",
        "articulo": "artículo",
        "logistica": "logística",
        "vehiculos": "vehículos",
        "automaticamente": "automáticamente",
        "autonomos": "autónomos",
        "asignacion": "asignación",
        "geometria": "geometría",
        "metodo": "método",
        "numero": "número",
        "teoria": "teoría",
        "optimizacion": "optimización",
        "saturacion": "saturación",
        "desempeno": "desempeño",
        "fisica": "física",
        "friccion": "fricción",
        "parametros": "parámetros",
        "matematica": "matemática",
        "minimo": "mínimo",
        "minima": "mínima",
        "codigo": "código",
        "diseno": "diseño",
        "Diseno": "Diseño",
        "Disenar": "Diseñar",
        "disenar": "diseñar",
        "implementacion": "implementación",
        "variacion": "variación",
        "maxima": "máxima",
        "aproximacion": "aproximación",
        "trayectoria": "trayectoria",
        "generacion": "generación",
        "graficas": "gráficas",
        "bibliografica": "bibliográfica",
        "revision": "revisión",
        "Revision": "Revisión",
        "decision": "decisión",
        "accion": "acción",
        "segun": "según",
        "mas": "más",
        "Mas": "Más",
        "esta": "está",
        "sera": "será",
        "podra": "podrá",
        "podran": "podrán",
        "mantendra": "mantendrá",
        "evaluara": "evaluará",
        "usara": "usará",
        "tendra": "tendrá",
        "trabajara": "trabajará",
        "realizara": "realizará",
        "separara": "separará",
        "incluiran": "incluirán",
        "quedara": "quedará",
        "podria": "podría",
        "podria": "podría",
        "tambien": "también",
        "despues": "después",
        "ultimos": "últimos",
        "practica": "práctica",
        "tecnico": "técnico",
        "tecnica": "técnica",
        "academica": "académica",
        "Academica": "Académica",
        "homogeneo": "homogéneo",
        "heterogeneo": "heterogéneo",
        "heterogeneas": "heterogéneas",
        "homogeneidad": "homogeneidad",
        "heterogeneidad": "heterogeneidad",
        "grafo": "grafo",
        "Nash": "Nash",
        "robotico": "robótico",
        "cinetico": "cinético",
        "cinematico": "cinemático",
        "cinematica": "cinemática",
        "cronograma": "cronograma",
    }
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        pattern = rf"(?<!\w){re.escape(source)}(?!\w)"
        text = re.sub(pattern, target, text)
    return text


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_paragraph_format(paragraph, *, after=8, before=0, line=1.333, justify=True) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def style_run(run, *, size=11, bold=False, italic=False, color=TEXT) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(fix_text(text))
        style_run(run)
    set_paragraph_format(paragraph)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(fix_text(text))
    if level == 1:
        style_run(run, size=16, bold=True, color=BLUE)
        set_paragraph_format(paragraph, before=18, after=10, line=1.2, justify=False)
    elif level == 2:
        style_run(run, size=13, bold=True, color=BLUE)
        set_paragraph_format(paragraph, before=12, after=6, line=1.2, justify=False)
    else:
        style_run(run, size=12, bold=True, color=DARK_BLUE)
        set_paragraph_format(paragraph, before=8, after=4, line=1.2, justify=False)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(fix_text(item))
        style_run(run)
        set_paragraph_format(p, after=4, line=1.208)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(fix_text(item))
        style_run(run)
        set_paragraph_format(p, after=4, line=1.208)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TABLE_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fix_text(header))
        style_run(r, size=10, bold=True, color=DARK_BLUE)
        set_paragraph_format(p, after=0, line=1.1, justify=False)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(fix_text(value))
            style_run(r, size=10)
            set_paragraph_format(p, after=0, line=1.15, justify=False)
    doc.add_paragraph()


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(fix_text(PROGRAM).upper())
    style_run(r, size=13, bold=True, color=DARK_BLUE)
    set_paragraph_format(p, after=20, line=1.2, justify=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(fix_text(TITLE))
    style_run(r, size=20, bold=True, color=BLUE)
    set_paragraph_format(p, after=12, line=1.15, justify=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(fix_text(SUBTITLE))
    style_run(r, size=12, italic=True, color=MUTED)
    set_paragraph_format(p, after=36, line=1.2, justify=False)

    meta = [
        f"Alumno: {AUTHOR}",
        f"Director: {TUTOR}",
        f"Institucion: {INSTITUTION}",
        "Edicion: 2025-26",
        f"Fecha: {DATE_TEXT}",
    ]
    for item in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fix_text(item))
        style_run(r, size=11, color=TEXT)
        set_paragraph_format(p, after=6, line=1.15, justify=False)

    doc.add_page_break()


def build_doc() -> None:
    if LATEX_SOURCE.exists():
        subprocess.run(
            [
                "pandoc",
                str(LATEX_SOURCE),
                "--from",
                "latex",
                "--to",
                "docx",
                "--number-sections",
                "--toc",
                "--metadata",
                f"title={TITLE}",
                "--metadata",
                f"author={AUTHOR}",
                "--metadata",
                f"date={DATE_TEXT}",
                "--metadata",
                "reference-section-title=Referencias bibliográficas",
                "--citeproc",
                "--bibliography",
                str(BIB_SOURCE),
                "--output",
                str(OUT_DOCX),
            ],
            cwd=ROOT,
            check=True,
        )
        doc = Document(OUT_DOCX)
        setup_document(doc)
        doc.save(OUT_DOCX)
        return

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc)

    add_heading(doc, "Resumen", 1)
    add_paragraph(
        doc,
        "Este Trabajo Fin de Master propone una arquitectura de control distribuido "
        "para transporte cooperativo multi-AGV en escenarios con cargas heterogeneas. "
        "El problema se formula como la necesidad de que una flota de robots moviles "
        "forme coaliciones de tamano adecuado para transportar cargas que pueden "
        "superar la capacidad individual de un agente, sin depender de un coordinador "
        "central con informacion completa.",
    )
    add_paragraph(
        doc,
        "La propuesta combina dos ideas. La primera es el control cooperativo basado "
        "en consenso, usado para mantener coordinacion local y movimiento hacia la "
        "carga asignada. La segunda es una dinamica poblacional distribuida, de tipo "
        "replicador regularizado o comparacion por pares, usada para ajustar en linea "
        "la preferencia de cada AGV por las cargas disponibles segun distancia, "
        "necesidad de coalicion, saturacion del grupo y calidad de comunicacion.",
    )
    add_paragraph(
        doc,
        "El alcance se limita a simulacion bidimensional reproducible en Python. La "
        "validacion principal compara el metodo propuesto con lineas base de consenso "
        "nominal, asignacion local voraz y asignacion centralizada ideal. La finalidad "
        "del TFM es cerrar una contribucion viable para master y, al mismo tiempo, "
        "dejar una linea clara hacia articulo de conferencia y posterior investigacion "
        "doctoral.",
    )
    add_paragraph(
        doc,
        "Palabras clave: dinamicas poblacionales, juegos evolutivos, control distribuido, "
        "formacion de coaliciones, consenso, AGV, transporte cooperativo.",
    )

    add_heading(doc, "Contenido", 1)
    add_numbered(
        doc,
        [
            "Introduccion",
            "Planteamiento del problema y justificacion",
            "Pregunta de investigacion e hipotesis",
            "Objetivos",
            "Alcance y limitaciones",
            "Metodologia",
            "Referencias bibliograficas",
        ],
    )

    add_heading(doc, "1. Introduccion", 1)
    add_paragraph(
        doc,
        "La automatizacion logistica y manufacturera depende cada vez mas de flotas "
        "de vehiculos guiados automaticamente y robots moviles autonomos. En muchos "
        "casos, estos sistemas trabajan bajo el supuesto de que cada robot transporta "
        "una carga individual de tamano y masa compatibles con su capacidad nominal. "
        "Ese supuesto simplifica la asignacion de tareas, pero pierde fuerza cuando "
        "aparecen cargas heterogeneas, piezas voluminosas o elementos cuya manipulacion "
        "requiere soporte distribuido.",
    )
    add_paragraph(
        doc,
        "El transporte cooperativo multi-robot aborda ese problema permitiendo que "
        "varios agentes actuen sobre una misma carga. La literatura reciente muestra "
        "avances en transporte cooperativo, control distribuido, esquemas basados en "
        "fuerza y control predictivo no lineal descentralizado (An et al., 2023; "
        "Ebel et al., 2024; Rosenfelder et al., 2024; Verginis et al., 2018). Estos "
        "trabajos demuestran que el problema sigue activo y que su dificultad no se "
        "reduce a mover robots, sino a coordinar decisiones, comunicacion, restricciones "
        "fisicas y desempeno colectivo.",
    )
    add_paragraph(
        doc,
        "El estado reciente de la literatura permite ubicar la brecha con mayor "
        "precision. Los grupos que lideran el transporte cooperativo fisico, como "
        "Ebel, Rosenfelder y Eberhard en Stuttgart, resuelven el control de formacion, "
        "fuerza y restricciones no holonomicas sobre el objeto transportado, pero no "
        "abordan como problema principal la seleccion distribuida de que robots atienden "
        "que carga cuando existen multiples cargas con cardinalidad variable. Los "
        "trabajos de asignacion distribuida para transporte colectivo, como Shan et al. "
        "(2024), usan subastas descentralizadas para tareas dinamicas con restricciones "
        "temporales, pero no integran la decision con una ley de control poblacional ni "
        "con una garantia formal de convergencia. Los enfoques por aprendizaje, como "
        "Shibata et al. (2023), Paul et al. (2023) y Bezerra et al. (2025), logran "
        "coordinacion y formacion de coaliciones en simulacion, aunque dependen de "
        "entrenamiento previo y ofrecen menos interpretabilidad formal. Por su parte, "
        "las dinamicas poblacionales de Quijano et al., Barreiro-Gomez et al. y "
        "Martinez-Piazuelo et al. han demostrado valor en asignacion de recursos y "
        "control de formacion, pero no se han usado directamente para formar coaliciones "
        "de cardinalidad variable en transporte cooperativo multi-AGV.",
    )
    add_paragraph(
        doc,
        "En paralelo, los juegos poblacionales y las dinamicas evolutivas han sido "
        "usados en control distribuido, optimizacion y formacion multiagente "
        "(Barreiro-Gomez et al., 2017; Barreiro-Gomez et al., 2021; Martinez-Piazuelo "
        "et al., 2020; Quijano et al., 2017; Sandholm, 2010). Esa linea resulta "
        "atractiva para este TFM porque permite describir la evolucion de preferencias "
        "locales sin exigir un planificador central. Tambien ofrece una base matematica "
        "para estudiar convergencia, asignacion de recursos y adaptacion bajo informacion "
        "parcial.",
    )
    add_paragraph(
        doc,
        "La oportunidad de investigacion esta en ese cruce especifico. "
        "Los metodos de asignacion multi-robot suelen separar decision y control: "
        "primero asignan robots a tareas y despues ejecutan trayectorias. En cambio, "
        "este TFM propone estudiar un ciclo acoplado donde la decision local, "
        "la formacion de coaliciones y el movimiento cooperativo evolucionan juntos. "
        "La contribucion no pretende "
        "resolver todos los casos de transporte cooperativo. Se concentra en una "
        "pregunta acotada: si una dinamica poblacional distribuida puede mejorar la "
        "formacion de coaliciones y el reparto de esfuerzo frente a baselines simples "
        "cuando la informacion es local y las cargas son heterogeneas.",
    )

    add_heading(doc, "2. Planteamiento del problema y justificacion", 1)
    add_paragraph(
        doc,
        "Se considera una flota de AGVs diferenciales que opera en un entorno plano. "
        "En el espacio de trabajo aparecen cargas con distinta masa, geometria "
        "simplificada y destino. Cada carga requiere un numero minimo de robots para "
        "ser transportada. Ese requisito se modela mediante una cardinalidad n_k, "
        "calculada a partir de la masa o dificultad relativa de la tarea. El problema "
        "consiste en que los AGVs formen coaliciones factibles, se desplacen hacia "
        "las cargas asignadas y contribuyan al transporte sin informacion global "
        "perfecta.",
    )
    add_paragraph(
        doc,
        "La comunicacion se representa mediante un grafo local variable. Dos robots "
        "solo intercambian informacion si estan dentro de un radio de comunicacion "
        "o si existe un enlace activo en el grafo. Por tanto, cada agente conoce una "
        "parte del estado colectivo y debe estimar de forma distribuida cuantos robots "
        "estan interesados en cada carga. Esta restriccion hace que una asignacion "
        "centralizada optima sea util como referencia, pero no como solucion principal.",
    )
    add_paragraph(
        doc,
        "La justificacion academica del trabajo es doble. Desde control distribuido, "
        "el problema conecta con consenso, grafos de comunicacion y coordinacion "
        "multiagente (Bullo et al., 2009; Olfati-Saber y Murray, 2004; Ren y Beard, "
        "2008). Desde teoria de juegos, conecta con dinamicas poblacionales aplicadas "
        "a optimizacion y control (Barreiro-Gomez et al., 2017; Quijano et al., 2017). "
        "Desde robotica cooperativa, se ubica en un campo donde los enfoques predictivos "
        "y de optimizacion distribuida son referencias fuertes, pero a menudo resultan "
        "costosos para una implementacion inicial de TFM.",
    )
    add_paragraph(
        doc,
        "El valor del TFM no depende de afirmar una primicia absoluta. La posicion "
        "defendible es mas precisa: hasta donde alcanza la revision inicial, las "
        "dinamicas poblacionales han sido exploradas en control distribuido y formacion, "
        "mientras que la formacion de coaliciones para transporte cooperativo suele "
        "tratarse con subastas (Choi et al., 2009; Shan et al., 2024), juegos hedonicos "
        "(Dutta et al., 2021), aprendizaje (Shibata et al., 2023; Paul et al., 2023; "
        "Bezerra et al., 2025) u optimizacion. Frente a Shan et al. (2024), este TFM "
        "no propone un protocolo de subasta, sino una regla de revision poblacional "
        "con payoff de umbral, acoplada al movimiento fisico y con una prueba formal "
        "minima en el caso homogeneo. Queda espacio para evaluar una arquitectura "
        "sencilla que use dinamicas poblacionales como mecanismo de asignacion "
        "adaptativa dentro de un sistema de transporte cooperativo con comunicacion local.",
    )
    add_table(
        doc,
        ["Linea", "Referencias clave", "Que resuelve", "Limite relevante para este TFM"],
        [
            [
                "Transporte cooperativo fisico",
                "Ebel et al. (2024); Rosenfelder et al. (2024)",
                "Organizacion y control del transporte no prensil con robots diferenciales, restricciones no holonomicas y experimentacion real.",
                "No aborda como problema principal la asignacion distribuida de robots a multiples cargas con cardinalidades heterogeneas.",
            ],
            [
                "Asignacion para transporte colectivo",
                "Shan et al. (2024)",
                "Asignacion descentralizada por subasta para tareas dinamicas de transporte colectivo con restricciones temporales.",
                "Se concentra en la planificacion/asignacion de alto nivel; no busca garantias de convergencia basadas en dinamicas poblacionales ni un control fisico detallado de la carga.",
            ],
            [
                "Coaliciones con aprendizaje",
                "Shibata et al. (2023); Paul et al. (2023); Bezerra et al. (2025)",
                "Uso de MARL, graph reinforcement learning o MAPPO para coordinacion, asignacion, escalabilidad y formacion dinamica de coaliciones.",
                "Requiere entrenamiento y validacion empirica; las garantias formales y la interpretabilidad de la ley distribuida son mas limitadas.",
            ],
            [
                "Dinamicas poblacionales y juegos",
                "Quijano et al. (2017); Barreiro-Gomez et al. (2017); Martinez-Piazuelo et al. (2020)",
                "Marco formal para control distribuido, asignacion de recursos, optimizacion y formacion multiagente.",
                "La aplicacion directa a transporte cooperativo con cargas de distinta cardinalidad sigue abierta en la revision inicial.",
            ],
            [
                "Propuesta del TFM",
                "Este trabajo",
                "Integra decision local, formacion de coaliciones por cardinalidad y movimiento cooperativo en simulacion reproducible.",
                "Acota la contribucion a simulacion 2D, con prueba formal solo para un caso homogeneo simplificado.",
            ],
        ],
        [1650, 2100, 2800, 2810],
    )
    add_paragraph(
        doc,
        "La brecha, por tanto, no es que falten metodos de transporte cooperativo ni "
        "metodos de asignacion multi-robot. La brecha es la integracion de una regla "
        "decisional distribuida, interpretable y sin entrenamiento previo con un ciclo "
        "de movimiento cooperativo donde las cargas tienen requisitos minimos de "
        "coalicion. Esa es la zona especifica donde se posiciona este TFM.",
    )

    add_heading(doc, "3. Pregunta de investigacion e hipotesis", 1)
    add_paragraph(
        doc,
        "Pregunta de investigacion. Puede una arquitectura distribuida que combine "
        "control por consenso y dinamicas poblacionales mejorar la formacion de "
        "coaliciones y el transporte cooperativo multi-AGV de cargas heterogeneas, "
        "manteniendo un costo computacional menor que una asignacion centralizada "
        "con informacion completa?",
    )
    add_paragraph(doc, "Hipotesis de trabajo:")
    add_numbered(
        doc,
        [
            "En escenarios con cargas heterogeneas, la dinamica poblacional distribuida formara coaliciones factibles con mayor frecuencia que una asignacion local voraz basada solo en distancia.",
            "El metodo propuesto tendra menor rendimiento que un asignador central ideal en algunos escenarios, pero mantendra una brecha aceptable si el grafo conserva conectividad temporal suficiente.",
            "La degradacion del radio de comunicacion afectara el tiempo de formacion de coalicion y el throughput, pero el uso de estimaciones locales por consenso reducira fallos abruptos de coordinacion.",
            "La integracion de dinamicas poblacionales y consenso ofrecera una base mas publicable que una comparacion aislada entre controladores nominales, porque permite estudiar asignacion, comunicacion y control dentro de una misma metodologia.",
        ],
    )
    add_paragraph(
        doc,
        "Las hipotesis se contrastaran por simulacion reproducible, usando los mismos "
        "escenarios, semillas y metricas para todos los metodos comparados.",
    )

    add_heading(doc, "4. Objetivos", 1)
    add_heading(doc, "4.1. Objetivo general", 2)
    add_paragraph(
        doc,
        "Disenar y evaluar en simulacion una arquitectura de control distribuido "
        "basada en dinamicas poblacionales y consenso para la formacion adaptativa "
        "de coaliciones en transporte cooperativo multi-AGV con cargas heterogeneas "
        "y comunicacion local.",
    )
    add_heading(doc, "4.2. Objetivos especificos", 2)
    add_numbered(
        doc,
        [
            "Formalizar un modelo bidimensional del sistema, incluyendo AGVs diferenciales, cargas heterogeneas, destinos, requisitos de cardinalidad y grafo de comunicacion local.",
            "Definir una linea base distribuida basada en consenso y una linea base local voraz para comparar el comportamiento del metodo propuesto.",
            "Disenar una funcion de payoff que combine necesidad de coalicion, distancia a la carga, saturacion del grupo, esfuerzo esperado y calidad de comunicacion.",
            "Implementar una dinamica poblacional discreta para actualizar las preferencias de cada AGV por las cargas disponibles sin usar informacion global completa.",
            "Integrar la decision poblacional con un controlador local de movimiento y formacion basado en consenso o campos potenciales simples.",
            "Construir un pipeline reproducible de simulacion en Python con configuraciones, semillas, metricas, tablas y figuras exportables.",
            "Evaluar el desempeno mediante escenarios de homogeneidad, heterogeneidad, degradacion de comunicacion, fallos de agentes y variacion de dificultad de carga.",
            "Preparar un borrador de articulo orientado a conferencia, por ejemplo CDC, ECC o NecSys, que contenga la formulacion del juego poblacional con payoff de umbral, la proposicion de convergencia del caso homogeneo y resultados de los escenarios B y C.",
        ],
    )

    add_heading(doc, "5. Alcance y limitaciones", 1)
    add_heading(doc, "5.1. Alcance", 2)
    add_paragraph(
        doc,
        "El alcance se define para una entrega de master, no para un sistema industrial "
        "cerrado. La validacion principal sera numerica y reproducible. Se trabajara "
        "con un entorno bidimensional, entre 4 y 8 AGVs en los escenarios principales, "
        "con extensiones de 15-20 AGVs en el escenario D para evaluar tendencias de "
        "escalabilidad bajo degradacion de comunicacion. En los escenarios principales "
        "se consideraran entre 2 y 4 cargas por episodio. Las cargas se modelaran por "
        "posicion, destino, dificultad relativa y cardinalidad minima requerida.",
    )
    add_bullets(
        doc,
        [
            "control local de movimiento con modelo unicycle discretizado;",
            "grafos de comunicacion por radio o matriz de adyacencia variable;",
            "dinamica poblacional discreta para asignacion distribuida de preferencias;",
            "control por consenso o campos potenciales para movimiento y cohesion local;",
            "comparacion contra baselines simples y contra un asignador central ideal;",
            "analisis cuantitativo con metricas definidas antes de ejecutar la campana experimental.",
        ],
    )
    add_heading(doc, "5.2. Limitaciones", 2)
    add_paragraph(
        doc,
        "Quedan fuera del alcance la implementacion en hardware real, la integracion "
        "obligatoria con ROS2, la prueba formal completa de convergencia para el caso "
        "heterogeneo con grafo variable, la planificacion con obstaculos complejos y "
        "la implementacion completa de un controlador NMPC o DMPC. CoppeliaSim podra "
        "usarse como demostrador visual si el cronograma lo permite, pero no sera "
        "criterio de exito de la entrega principal.",
    )
    add_paragraph(
        doc,
        "Esta delimitacion protege la viabilidad del TFM. Las lineas que quedan fuera "
        "del alcance constituyen la ruta doctoral prevista: demostracion formal de "
        "convergencia para el caso heterogeneo con grafo variable, usando separacion "
        "de escalas temporales y perturbaciones singulares; extension hacia control "
        "estigmergico, donde la comunicacion explicita se complemente o reemplace por "
        "marcadores ambientales; validacion experimental con robots reales; y comparacion "
        "formal contra enfoques NMPC distribuidos, subastas descentralizadas y MARL.",
    )

    add_heading(doc, "6. Metodologia", 1)
    add_paragraph(
        doc,
        "La investigacion sera aplicada, cuantitativa y experimental por simulacion. "
        "El trabajo seguira un ciclo de modelado, diseno de control, implementacion, "
        "experimentacion y analisis. La metodologia prioriza comparaciones trazables "
        "sobre promesas de generalidad. Cada resultado debera poder reproducirse a "
        "partir de configuraciones versionadas, semillas conocidas y scripts de "
        "generacion de metricas.",
    )

    add_heading(doc, "6.1. Fase 1: revision focalizada y modelo formal", 2)
    add_paragraph(
        doc,
        "La primera fase cerrara la formulacion matematica minima. Cada AGV i tendra "
        "estado q_i = (x_i, y_i, theta_i) y entradas de velocidad lineal y angular. "
        "Las cargas k tendran posicion, destino, recompensa o prioridad, dificultad "
        "relativa y cardinalidad n_k. El grafo G(t) describira que pares de AGVs pueden "
        "intercambiar informacion en cada instante.",
    )
    add_paragraph(
        doc,
        "La variable de decision principal sera una distribucion de preferencias "
        "p_i(t) sobre las cargas y, si conviene, sobre una estrategia idle. Esta "
        "representacion permite que cada robot no elija de forma binaria desde el "
        "primer instante, sino que actualice gradualmente su inclinacion hacia cada "
        "tarea en funcion de los payoffs locales.",
    )

    add_heading(doc, "6.2. Fase 2: lineas base", 2)
    add_paragraph(
        doc,
        "Se implementaran tres referencias de dificultad creciente. La primera sera "
        "una asignacion local voraz, donde cada AGV elige la carga mas cercana que "
        "todavia parece necesitar robots. La segunda sera una variante distribuida "
        "basada en consenso para estimar demanda y coordinar movimiento. La tercera "
        "sera un asignador central ideal que usa informacion global para construir "
        "coaliciones factibles. Esta ultima no representa una solucion distribuida, "
        "sino un techo de rendimiento para medir la brecha de descentralizacion.",
    )

    add_heading(doc, "6.3. Fase 3: metodo propuesto", 2)
    add_paragraph(
        doc,
        "El metodo propuesto asignara a cada robot una funcion de payoff local para "
        "cada carga. La forma principal sera multiplicativa: f_ik = rho_k Phi(z_ik, n_k) "
        "Psi(d_ik). En esta expresion, rho_k representa prioridad de la carga, Phi mide "
        "demanda remanente de coalicion y Psi mide accesibilidad espacial.",
    )
    add_paragraph(
        doc,
        "La funcion de demanda se definira como Phi(z, n) = 1 / (1 + exp(beta (z - n))). "
        "Asi, una carga cercana a su cardinalidad requerida reduce suavemente su atractivo. "
        "Si se desea que cargas con mayor cardinalidad sean intrinsecamente mas atractivas, "
        "ese efecto se incorporara en rho_k, no duplicando la funcion de demanda.",
    )
    add_paragraph(
        doc,
        "Las preferencias se actualizaran mediante una dinamica poblacional discreta. "
        "La version principal usara una dinamica de Smith discreta con proyeccion al "
        "simplex, porque permite que una estrategia con preferencia momentaneamente "
        "nula pueda recuperarse si su payoff vuelve a ser competitivo.",
    )
    add_paragraph(
        doc,
        "La salida de la capa poblacional no sustituira al controlador cinematico. "
        "Servira para decidir a que carga tiende cada robot y con que peso participa "
        "en la coalicion. El movimiento se resolvera con un controlador local simple "
        "basado en atraccion hacia la carga asignada, repulsion entre AGVs cercanos "
        "y consenso para mantener cohesion entre robots que comparten una misma carga.",
    )

    add_heading(doc, "6.4. Fase 4: implementacion reproducible", 2)
    add_paragraph(
        doc,
        "La implementacion se realizara en Python. Cada experimento tendra un archivo "
        "de configuracion con numero de AGVs, numero de cargas, posiciones iniciales, "
        "cardinalidades, radio de comunicacion, semillas, parametros de payoff y "
        "ganancias de control. Los resultados brutos se guardaran separados de las "
        "tablas procesadas y de las figuras finales.",
    )
    add_paragraph(
        doc,
        "La arquitectura del repositorio separara dominio, controladores, dinamicas "
        "poblacionales, simulacion, metricas, experimentos y visualizacion. Esta "
        "separacion es importante porque el TFM debe poder crecer despues hacia un "
        "paper sin reescribir todo el codigo experimental.",
    )

    add_heading(doc, "6.5. Fase 5: diseno experimental", 2)
    add_paragraph(doc, "La campana experimental se organizara en cinco escenarios.")
    add_table(
        doc,
        ["Escenario", "Proposito", "Resultado esperado"],
        [
            [
                "A. Homogeneo",
                "Cargas con cardinalidad igual o similar.",
                "Verificar que el sistema no introduce conflictos innecesarios.",
            ],
            [
                "B. Coalicion simple",
                "Una carga ligera y una carga que exige varios AGVs.",
                "Comprobar si emerge una coalicion factible para la carga dificil.",
            ],
            [
                "C. Heterogeneo",
                "Varias cargas con cardinalidades distintas.",
                "Medir competencia entre tareas y reparto de la flota.",
            ],
            [
                "D. Comunicacion degradada",
                "Barrido del radio de comunicacion.",
                "Obtener la curva de degradacion frente a perdida de informacion local.",
            ],
            [
                "E. Fallo o incertidumbre",
                "Retiro de un AGV, cambio de dificultad o perturbacion acotada.",
                "Evaluar recuperacion de coaliciones y perdida relativa de rendimiento.",
            ],
        ],
        [1650, 3850, 3860],
    )
    add_paragraph(
        doc,
        "Cada escenario se ejecutara con varias semillas. El informe no se limitara "
        "a una trayectoria visual. Se reportaran metricas agregadas y, cuando sea "
        "necesario, intervalos de variacion para mostrar sensibilidad parametrica.",
    )

    add_heading(doc, "6.6. Fase 6: metricas y analisis", 2)
    add_table(
        doc,
        ["Metrica", "Definicion operativa"],
        [
            ["Tasa de coalicion factible", "Porcentaje de cargas atendidas con al menos n_k robots asignados."],
            ["Tiempo de formacion", "Tiempo hasta que una coalicion alcanza cardinalidad requerida y permanece estable durante una ventana minima."],
            ["Throughput", "Numero de cargas entregadas o completadas por unidad de tiempo simulado."],
            ["Error de seguimiento", "Distancia media o maxima entre la trayectoria de la coalicion y la trayectoria deseada de la carga."],
            ["Esfuerzo de control", "Magnitud acumulada de velocidades o comandos aplicados por los AGVs."],
            ["Brecha de descentralizacion", "Diferencia relativa entre el metodo distribuido y el baseline central ideal."],
            ["Costo computacional", "Tiempo medio de calculo por paso de simulacion y por agente."],
        ],
        [2700, 6660],
    )
    add_paragraph(
        doc,
        "El analisis comparara los metodos por escenario y por metrica. La superioridad "
        "del metodo propuesto no se definira como ganar en todas las metricas, sino "
        "como mejorar la factibilidad y adaptacion distribuida con una brecha razonable "
        "respecto al asignador central ideal. Tambien se reportaran casos negativos, "
        "por ejemplo radios de comunicacion donde la informacion local ya no basta "
        "para sostener coaliciones estables.",
    )

    add_heading(doc, "6.7. Control de validez y riesgos", 2)
    add_paragraph(
        doc,
        "La principal amenaza a la validez es que el simulador simplifique demasiado "
        "la fisica del transporte. Para mitigarlo, el TFM separara con claridad lo "
        "que se demuestra sobre asignacion y coordinacion de lo que queda pendiente "
        "sobre contacto, friccion y dinamica completa de la carga. Otra amenaza es "
        "la sensibilidad a parametros de payoff. Por eso se incluiran barridos sobre "
        "ganancias principales y no solo un ajuste manual favorable.",
    )
    add_paragraph(
        doc,
        "La tercera amenaza es sobreafirmar novedad. El documento evitara declarar "
        "que el metodo es el primero de su tipo. La formulacion adecuada es que la "
        "revision inicial no encontro una integracion directa entre dinamicas "
        "poblacionales distribuidas, formacion de coaliciones por cardinalidad y "
        "transporte cooperativo multi-AGV bajo comunicacion local. Esa afirmacion "
        "podra fortalecerse o corregirse durante la revision bibliografica completa.",
    )

    add_heading(doc, "Referencias bibliograficas", 1)
    references = [
        "An, X., Wu, C., Lin, Y., Lin, M., Yoshinaga, T. y Ji, Y. (2023). Multi-Robot Systems and Cooperative Object Transport: Communications, Platforms, and Challenges. IEEE Open Journal of the Computer Society, 4, 23-36.",
        "Barreiro-Gomez, J., Obando, G. y Quijano, N. (2017). Distributed population dynamics: Optimization and control applications. IEEE Transactions on Systems, Man, and Cybernetics: Systems, 47(2), 304-314.",
        "Barreiro-Gomez, J., Mas, I., Giribet, J. I., Moreno, P., Ocampo-Martinez, C., Sanchez-Pena, R. y Quijano, N. (2021). Distributed data-driven UAV formation control via evolutionary games: Experimental results. Journal of the Franklin Institute, 358(10), 5334-5352.",
        "Bezerra, L. C. D., dos Santos, A. M. G. y Park, S. (2025). Learning policies for dynamic coalition formation in multi-robot task allocation. IEEE Robotics and Automation Letters, 10(9), 9216-9223. DOI: 10.1109/LRA.2025.3592080.",
        "Bullo, F., Cortes, J. y Martinez, S. (2009). Distributed Control of Robotic Networks: A Mathematical Approach to Motion Coordination Algorithms. Princeton University Press.",
        "Choi, H.-L., Brunet, L. y How, J. P. (2009). Consensus-based decentralized auctions for robust task allocation. IEEE Transactions on Robotics, 25(4), 912-926.",
        "Dutta, A., Ufimtsev, V., Said, T., Jang, I. y Eggen, R. (2021). Distributed hedonic coalition formation for multi-robot task allocation. IEEE 17th International Conference on Automation Science and Engineering, 639-644.",
        "Ebel, H., Rosenfelder, M. y Eberhard, P. (2024). Cooperative object transportation with differential-drive mobile robots: Control and experimentation. Robotics and Autonomous Systems, 173, 104612. DOI: 10.1016/j.robot.2023.104612.",
        "Martinez-Piazuelo, J., Diaz-Garcia, G., Quijano, N. y Giraldo, L. F. (2020). Distributed formation control of mobile robots using discrete-time distributed population dynamics. IFAC-PapersOnLine, 53(2), 3131-3136.",
        "Olfati-Saber, R. y Murray, R. M. (2004). Consensus problems in networks of agents with switching topology and time-delays. IEEE Transactions on Automatic Control, 49(9), 1520-1533.",
        "Quijano, N., Ocampo-Martinez, C., Barreiro-Gomez, J., Obando, G., Pantoja, A. y Mojica-Nava, E. (2017). The role of population games and evolutionary dynamics in distributed control systems. IEEE Control Systems Magazine, 37(1), 70-97.",
        "Ren, W. y Beard, R. W. (2008). Distributed Consensus in Multi-Vehicle Cooperative Control: Theory and Applications. Springer.",
        "Paul, S., Li, W., Smyth, B., Chen, Y., Gel, Y. R. y Chowdhury, S. (2023). Efficient planning of multi-robot collective transport using graph reinforcement learning with higher order topological abstraction. 2023 IEEE International Conference on Robotics and Automation (ICRA), 5779-5785. DOI: 10.1109/ICRA48891.2023.10161517.",
        "Rosenfelder, M., Ebel, H. y Eberhard, P. (2024). Force-based organization and control scheme for the non-prehensile cooperative transportation of objects. Robotica, 42(2), 611-624. DOI: 10.1017/S0263574723001704.",
        "Sandholm, W. H. (2010). Population Games and Evolutionary Dynamics. MIT Press.",
        "Shan, X., Jin, Y., Jurt, M. y Li, P. (2024). A distributed multi-robot task allocation method for time-constrained dynamic collective transport. Robotics and Autonomous Systems, 178, 104722. DOI: 10.1016/j.robot.2024.104722.",
        "Shibata, K., Jimbo, T. y Matsubara, T. (2023). Deep reinforcement learning of event-triggered communication and consensus-based control for distributed cooperative transport. Robotics and Autonomous Systems, 159, 104307. DOI: 10.1016/j.robot.2022.104307.",
        "Shibata, K., Jimbo, T., Odashima, T., Takeshita, K. y Matsubara, T. (2023). Learning locally, communicating globally: Reinforcement learning of multi-robot task allocation for cooperative transport. IFAC-PapersOnLine, 56(2), 11436-11443. DOI: 10.1016/j.ifacol.2023.10.431.",
        "Verginis, C. K., Nikou, A. y Dimarogonas, D. V. (2018). Communication-based decentralized cooperative object transportation using nonlinear model predictive control. European Control Conference, 733-738.",
        "Wurman, P. R., D'Andrea, R. y Mountz, M. (2008). Coordinating hundreds of cooperative, autonomous vehicles in warehouses. AI Magazine, 29(1), 9-20.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        r = p.add_run(fix_text(ref))
        style_run(r, size=10)
        set_paragraph_format(p, after=6, line=1.15)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)

    # Footer with document label.
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(fix_text("TFM Mayorga Taborda - propuesta hasta metodologia"))
        style_run(run, size=9, color=MUTED)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOCX)
