"""Build the VIU MROB TFM Task 1 DOCX draft from the official template."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from lxml import etree
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\walla\Downloads\Plantilla memoria TFM_ MROB.docx")
OUT_DIR = ROOT / "docs" / "thesis-final"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "TFM_Tarea1_Fundamentacion_MROB_AGV.docx"

TITLE = (
    "Coordinación distribuida local de múltiples AGV para el transporte "
    "cooperativo de cargas heterogéneas en entornos industriales"
)
AUTHOR = "Jorge Luis Mayorga Taborda"
TUTOR = "Jose Ignacio Iñiguez Amigot"
DATE_TEXT = "20 de mayo de 2026"
EDITION = "2025-26"


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_normal(paragraph, size=12):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)


def add_p(doc: Document, text: str, style: str = "Normal", align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    set_normal(p)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(f"• {item}", style="Normal")
        set_normal(p)


def add_numbered(doc: Document, items: list[str]):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(f"{idx}. {item}", style="Normal")
        set_normal(p)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]):
    p = add_p(doc, title, "Normal")
    p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Table Grid", "Cuadrícula de tabla", "Tabla con cuadrícula"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    hdr_cells = table.rows[0].cells
    for cell, header in zip(hdr_cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
    doc.add_paragraph("")
    return table


def make_architecture_figure(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1900, 1040
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 34)
        small_font = ImageFont.truetype("arial.ttf", 27)
    except OSError:
        font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    def draw_multiline_box(text: str, center: tuple[int, int], box_size: tuple[int, int]):
        cx, cy = center
        bw, bh = box_size
        x0, y0 = cx - bw // 2, cy - bh // 2
        x1, y1 = cx + bw // 2, cy + bh // 2
        draw.rounded_rectangle((x0, y0, x1, y1), radius=28, fill="#f3f6fb", outline="#2f5597", width=4)
        lines = text.split("\n")
        line_h = 42
        total_h = line_h * len(lines)
        for i, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            draw.text((cx - tw / 2, cy - total_h / 2 + i * line_h), line, fill="#1f1f1f", font=font)

    def arrow(start: tuple[int, int], end: tuple[int, int]):
        draw.line((start, end), fill="#555555", width=5)
        ex, ey = end
        sx, sy = start
        dx, dy = ex - sx, ey - sy
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        size = 28
        p1 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
        p2 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
        draw.polygon((end, p1, p2), fill="#555555")

    boxes = [
        ("Problema industrial\nBodega automatizada\nCargas heterogéneas", (310, 315), (500, 210)),
        ("Teoría central\nControl cooperativo\nConsenso + grafos", (950, 315), (500, 210)),
        ("Método propuesto\nControl distribuido\nAdaptativo/robusto", (1590, 315), (500, 210)),
        ("Python\nSimulación reproducible\nMétricas y tablas", (650, 760), (500, 210)),
        ("CoppeliaSim\nValidación aplicada\nEscena multi-AGV", (1250, 760), (500, 210)),
    ]
    for text, center, size in boxes:
        draw_multiline_box(text, center, size)
    for start, end in [
        ((560, 315), (700, 315)),
        ((1200, 315), (1340, 315)),
        ((1530, 425), (770, 655)),
        ((1590, 425), (1320, 655)),
        ((950, 425), (690, 655)),
    ]:
        arrow(start, end)
    footer = "La teoría se centra en control y cooperación; Python y CoppeliaSim actúan como plataformas de contraste."
    bbox = draw.textbbox((0, 0), footer, font=small_font)
    draw.text(((width - (bbox[2] - bbox[0])) / 2, 970), footer, fill="#333333", font=small_font)
    image.save(path)


def strip_template_toc_field(docx_path: Path):
    """Remove the original template TOC field that python-docx does not expose."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path, "r") as src:
        files = {name: src.read(name) for name in src.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    removed = 0
    for sdt in list(root.xpath(".//w:sdt", namespaces=ns)):
        text = "".join(sdt.xpath(".//w:t/text() | .//w:instrText/text()", namespaces=ns))
        if 'TOC \\o "1-3"' in text or "Marco teórico (normativo)" in text:
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)
                removed += 1
    if removed:
        files["word/document.xml"] = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone="yes",
        )
        temp_path = docx_path.with_suffix(".tmp.docx")
        with ZipFile(temp_path, "w", ZIP_DEFLATED) as dst:
            for name, data in files.items():
                dst.writestr(name, data)
        temp_path.replace(docx_path)


def build_doc():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig_path = FIG_DIR / "tarea1_arquitectura_metodologica.png"
    make_architecture_figure(fig_path)

    doc = Document(TEMPLATE)

    # Fill cover placeholders.
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "TÍTULO:":
            p.text = f"TÍTULO:\n{TITLE}"
            set_normal(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
        elif text == "Alumno/a:":
            p.text = f"Alumno/a: {AUTHOR}"
            set_normal(p)
        elif text == "Director/a:":
            p.text = f"Director/a: {TUTOR}"
            set_normal(p)
        elif text.startswith("Edición:"):
            p.text = f"Edición: {EDITION}\nFecha: {DATE_TEXT}"
            set_normal(p)

    # Remove instructional body from "Resumen" onward.
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Resumen":
            start = i
            break
    if start is not None:
        for p in list(doc.paragraphs[start:]):
            remove_paragraph(p)

    doc.add_page_break()

    add_p(doc, "Resumen", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    add_p(
        doc,
        (
            "Este Trabajo Fin de Máster plantea el estudio de la coordinación distribuida local de múltiples "
            "AGV para el transporte cooperativo de cargas heterogéneas en entornos industriales. El problema "
            "se ubica en bodegas automatizadas donde una sola unidad móvil puede resultar insuficiente para "
            "desplazar cargas voluminosas, pesadas o con propiedades inerciales variables, y donde el uso de "
            "un control centralizado puede limitar la escalabilidad y la robustez del sistema. El objetivo del "
            "trabajo es diseñar, implementar y evaluar una arquitectura de control cooperativo distribuido, "
            "basada en consenso y con compensación adaptativa o robusta, que permita coordinar varios AGV a "
            "partir de información local. La metodología propuesta combina revisión bibliográfica, modelado "
            "planar del sistema robot-carga, definición de grafos de comunicación, diseño de una estrategia "
            "nominal y una estrategia adaptativa, simulación cuantitativa en Python y validación aplicada en "
            "CoppeliaSim. La comparación se realizará mediante métricas de error de seguimiento, error de "
            "formación, esfuerzo de control, tiempo de convergencia y robustez cooperativa frente a variaciones "
            "de masa, desplazamientos del centro de masas, degradación de comunicación y perturbaciones "
            "externas. Como resultado esperado, se busca obtener una evaluación técnica reproducible que "
            "permita identificar las condiciones bajo las cuales el control cooperativo adaptativo mejora el "
            "desempeño frente a una estrategia nominal."
        ),
    )
    p = add_p(
        doc,
        "Palabras clave: control distribuido; cooperación multiagente; AGV; transporte cooperativo; CoppeliaSim.",
    )
    p.runs[0].bold = True

    doc.add_page_break()
    add_p(doc, "Contenido", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    toc_items = [
        "1. Introducción",
        "1.1. Contexto industrial y problema de investigación",
        "1.2. Soluciones actuales, límites y brecha de investigación",
        "1.3. Plataforma propuesta y delimitación inicial",
        "1.4. Pregunta de investigación y justificación académica",
        "2. Objetivos",
        "2.1. Objetivo general",
        "2.2. Objetivos específicos",
        "2.3. Alcance del TFM",
        "3. Hipótesis de partida",
        "4. Metodología",
        "4.1. Diseño de investigación",
        "4.2. Fases del trabajo",
        "4.3. Plataforma de simulación y validación",
        "4.4. Método propuesto",
        "4.5. Metodología de comparación",
        "4.6. Variables, escenarios y métricas",
        "4.7. Limitaciones metodológicas",
        "4.8. Criterios de aceptación y trazabilidad",
        "Referencias bibliográficas",
    ]
    for item in toc_items:
        add_p(doc, item)

    add_p(doc, "Índice de figuras", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Figura 1. Arquitectura metodológica del TFM.")
    add_p(doc, "Índice de tablas", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    for item in [
        "Tabla 1. Relación entre problema, brecha y respuesta metodológica.",
        "Tabla 2. Variables de estudio.",
        "Tabla 3. Escenarios de evaluación.",
        "Tabla 4. Métricas de comparación.",
    ]:
        add_p(doc, item)

    doc.add_page_break()
    add_p(doc, "1. Introducción", "Heading 1")
    add_p(
        doc,
        (
            "La automatización de bodegas y centros de distribución ha incrementado el uso de vehículos "
            "guiados automáticamente y robots móviles para tareas de transporte interno. En sistemas de alta "
            "rotación, estos robots permiten reducir desplazamientos manuales, mejorar la repetibilidad de las "
            "operaciones y conectar estaciones de almacenamiento, preparación y despacho. Sin embargo, una "
            "parte importante de las soluciones industriales se diseña bajo una hipótesis operativa sencilla: "
            "cada robot transporta una unidad de carga compatible con sus dimensiones, capacidad mecánica y "
            "modo de sujeción. Esta hipótesis funciona bien cuando los objetos son homogéneos, pero resulta "
            "menos flexible cuando aparecen cargas heterogéneas, voluminosas, frágiles o demasiado pesadas "
            "para un único AGV."
        ),
    )
    add_p(
        doc,
        (
            "El transporte cooperativo multi-robot aborda precisamente ese tipo de situaciones. En lugar de "
            "asignar una carga a un único vehículo, varios agentes móviles actúan de forma coordinada para "
            "desplazar un mismo objeto o para mantener una formación que permita soportar y guiar la carga. "
            "Desde el punto de vista de la robótica avanzada, el problema no consiste únicamente en desplazar "
            "robots en el plano, sino en diseñar reglas de interacción que permitan que un comportamiento "
            "colectivo emerja a partir de decisiones locales. Esta idea conecta directamente con los sistemas "
            "multiagente, la teoría de grafos, el control distribuido y los algoritmos de consenso."
        ),
    )
    add_p(
        doc,
        (
            "El problema se vuelve más exigente cuando la carga transportada no se conoce con exactitud. En "
            "una bodega real, una misma plataforma puede manipular objetos con distintas masas, geometrías o "
            "distribuciones de peso. La incertidumbre en masa, inercia o centro de masas modifica la respuesta "
            "dinámica del sistema y puede degradar tanto el seguimiento de trayectoria como el mantenimiento "
            "de la formación entre AGV. Si además se impone una comunicación local, cada agente dispone solo "
            "de información parcial de sus vecinos y no de una visión completa del sistema. Este escenario "
            "motiva el uso de control cooperativo distribuido frente a soluciones centralizadas."
        ),
    )

    add_p(doc, "1.1. Contexto industrial y problema de investigación", "Heading 2")
    add_p(
        doc,
        (
            "En entornos industriales, los AGV y AMR suelen integrarse en flujos de producción, almacenamiento "
            "o intralogística. Las soluciones comerciales de gran escala demuestran que la coordinación de "
            "flotas puede ser altamente eficiente cuando las tareas son repetitivas y las cargas están "
            "normalizadas. No obstante, la cooperación física sobre una misma carga plantea una dificultad "
            "distinta: varios robots deben mantener relaciones geométricas, distribuir esfuerzos y responder "
            "a perturbaciones de forma coherente. En esta situación, el problema técnico principal no es solo "
            "la navegación individual, sino la coordinación dinámica del conjunto."
        ),
    )
    add_p(
        doc,
        (
            "El problema de investigación de este TFM se formula así: cómo diseñar una estrategia de control "
            "cooperativo distribuido que permita a varios AGV transportar cargas heterogéneas bajo comunicación "
            "local y parámetros inerciales inciertos. La formulación se centra en el plano, con robots móviles "
            "de tracción diferencial o comportamiento cinemático equivalente, y con una carga rígida cuya masa "
            "o distribución inercial puede variar entre escenarios. La decisión de trabajar en el plano acota "
            "el problema y permite concentrar el esfuerzo en el diseño de control y cooperación."
        ),
    )

    add_p(doc, "1.2. Soluciones actuales, límites y brecha de investigación", "Heading 2")
    add_p(
        doc,
        (
            "Las soluciones existentes pueden agruparse en tres familias. La primera corresponde a esquemas "
            "centralizados, donde una unidad de planificación o control calcula las acciones del conjunto. "
            "Estos enfoques simplifican la coordinación global, pero pueden crear dependencia de un nodo "
            "central, aumentar los requisitos de comunicación y reducir la tolerancia a fallos. La segunda "
            "familia corresponde a arquitecturas líder-seguidor, en las que un robot o una referencia virtual "
            "organiza el movimiento de los demás. Esta solución reduce parte de la complejidad, aunque mantiene "
            "dependencia de un rol privilegiado y puede ser sensible a la pérdida o degradación del líder. La "
            "tercera familia corresponde a control distribuido y consenso, donde cada agente actualiza su "
            "acción a partir de información local y del estado relativo de sus vecinos."
        ),
    )
    add_p(
        doc,
        (
            "La literatura de consenso y cooperación multiagente muestra que es posible coordinar vehículos "
            "sin requerir comunicación global, siempre que la topología de comunicación satisfaga condiciones "
            "mínimas de conectividad (Olfati-Saber et al., 2007; Ren & Beard, 2008). Sin embargo, en transporte "
            "cooperativo de cargas heterogéneas aparecen límites adicionales: la carga introduce acoplamiento "
            "físico entre agentes, los parámetros inerciales pueden ser inciertos y la degradación de la "
            "comunicación puede afectar directamente la formación. Por tanto, la brecha que aborda este TFM "
            "no es demostrar que el consenso existe, sino evaluar cómo una arquitectura cooperativa distribuida "
            "se comporta cuando el transporte de carga introduce incertidumbre dinámica y escenarios de "
            "operación más cercanos a una bodega industrial."
        ),
    )
    add_p(
        doc,
        (
            "En el contexto industrial, las plataformas de bodega automatizada han demostrado que la coordinación "
            "de grandes flotas es viable cuando las unidades de carga, las estanterías y los flujos de trabajo "
            "están estandarizados. No obstante, esa madurez no resuelve completamente el caso de cargas "
            "heterogéneas que requieren cooperación física simultánea. En el primer caso, cada robot ejecuta "
            "una tarea individual dentro de una flota coordinada; en el segundo, varios robots deben actuar "
            "como un sistema mecánico distribuido. Esta diferencia justifica que el TFM se centre en cooperación "
            "y control, y no únicamente en planificación logística o gestión de flotas."
        ),
    )
    add_p(
        doc,
        (
            "Desde el punto de vista académico, el trabajo se apoya en un conjunto de antecedentes consolidados: "
            "la teoría de consenso para sistemas multiagente, la coordinación de redes robóticas, la cooperación "
            "en sistemas de robots móviles y el transporte cooperativo de objetos. Sin embargo, el interés del "
            "TFM está en articular esos elementos en un marco aplicado a AGV industriales con comunicación local "
            "y carga incierta. Por tanto, el aporte no se define como la invención de un principio de control "
            "completamente nuevo, sino como la formulación, implementación y evaluación rigurosa de una "
            "arquitectura cooperativa bajo condiciones de operación delimitadas."
        ),
    )

    add_table(
        doc,
        "Tabla 1. Relación entre problema, brecha y respuesta metodológica.",
        ["Elemento", "Planteamiento en el TFM"],
        [
            [
                "Problema técnico",
                "Transporte cooperativo de cargas heterogéneas con múltiples AGV bajo incertidumbre inercial.",
            ],
            [
                "Soluciones actuales",
                "Control centralizado, líder-seguidor, consenso distribuido y enfoques de cooperación multi-robot.",
            ],
            [
                "Límites detectados",
                "Dependencia de comunicación global, sensibilidad a parámetros de carga y evaluación limitada ante degradación del grafo.",
            ],
            [
                "Respuesta propuesta",
                "Arquitectura de control cooperativo distribuido con comparación nominal/adaptativa en Python y CoppeliaSim.",
            ],
        ],
    )

    add_p(doc, "1.3. Plataforma propuesta y delimitación inicial", "Heading 2")
    add_p(
        doc,
        (
            "El trabajo plantea dos niveles de plataforma. El primer nivel es una simulación cuantitativa en "
            "Python, donde se implementan el modelo, los controladores, las perturbaciones, las métricas y la "
            "generación de resultados. Este nivel será la base científica del trabajo porque permite controlar "
            "condiciones iniciales, repetir escenarios y comparar de forma justa estrategias de control. El "
            "segundo nivel es CoppeliaSim, usado como entorno de validación aplicada para representar la escena "
            "multi-AGV en una bodega, visualizar la cooperación y comprobar que el enfoque no queda reducido a "
            "una simulación abstracta sin conexión robótica."
        ),
    )
    add_p(
        doc,
        (
            "La delimitación del TFM excluye la construcción física de AGV, la percepción real, la planificación "
            "global de rutas, la navegación SLAM y la certificación de seguridad industrial. Estas decisiones "
            "no reducen el valor del trabajo, sino que permiten focalizarlo en el núcleo académico: control y "
            "cooperación. La validación se realizará mediante escenarios simulados, con especial atención al "
            "seguimiento de trayectoria, mantenimiento de formación, esfuerzo de control, tiempo de convergencia "
            "y robustez frente a incertidumbre."
        ),
    )

    add_p(doc, "1.4. Pregunta de investigación y justificación académica", "Heading 2")
    add_p(
        doc,
        (
            "La pregunta de investigación que orienta el trabajo es la siguiente: ¿puede una arquitectura de "
            "control cooperativo distribuido, basada en consenso y con compensación adaptativa o robusta, "
            "mejorar el transporte de cargas heterogéneas por múltiples AGV bajo incertidumbre inercial y "
            "comunicación local, en comparación con una estrategia nominal? Esta pregunta es adecuada para un "
            "TFM de robótica y automatización porque integra modelado, control, cooperación, simulación y "
            "validación aplicada."
        ),
    )
    add_p(
        doc,
        (
            "La justificación académica se apoya en tres argumentos. Primero, el problema pertenece al ámbito de "
            "la robótica móvil industrial y de los sistemas multiagente, áreas directamente vinculadas con el "
            "Máster Universitario en Robótica y Automatización de Procesos. Segundo, el trabajo exige aplicar "
            "conocimientos de control avanzado, modelado de sistemas y simulación robótica. Tercero, la "
            "evaluación comparativa permite producir evidencia cuantitativa, no solo una demostración visual. "
            "Con ello, el documento puede cumplir el doble carácter esperado: memoria técnica y estudio "
            "teórico-técnico."
        ),
    )
    add_p(
        doc,
        (
            "La memoria se organizará de forma progresiva. En primer lugar, se delimita el problema y se ubican "
            "las soluciones actuales. En segundo lugar, se definen objetivos, hipótesis y alcance. En tercer "
            "lugar, se desarrolla la metodología de modelado, control, simulación y comparación. En la versión "
            "final se incorporarán el marco teórico completo, los resultados, el análisis, la validación y las "
            "conclusiones."
        ),
    )
    doc.add_picture(str(fig_path), width=Inches(6.2))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, "Figura 1. Arquitectura metodológica del TFM.", "Normal", WD_ALIGN_PARAGRAPH.CENTER)

    add_p(doc, "2. Objetivos", "Heading 1")
    add_p(
        doc,
        (
            "Los objetivos se formulan para que el trabajo sea evaluable mediante resultados técnicos. El "
            "objetivo general expresa el propósito principal, mientras que los objetivos específicos organizan "
            "el desarrollo en tareas verificables: modelado, cooperación, control, simulación, validación y "
            "comparación."
        ),
    )
    add_p(doc, "2.1. Objetivo general", "Heading 2")
    add_p(
        doc,
        (
            "Diseñar, implementar y evaluar una arquitectura de control cooperativo distribuido para múltiples "
            "AGV, basada en consenso y con compensación adaptativa o robusta, orientada al transporte de cargas "
            "heterogéneas bajo incertidumbre inercial en entornos industriales simulados."
        ),
    )
    add_p(doc, "2.2. Objetivos específicos", "Heading 2")
    add_numbered(
        doc,
        [
            "Caracterizar el problema de transporte cooperativo de cargas heterogéneas en una bodega automatizada.",
            "Formular un modelo planar del sistema multi-AGV-carga con variables cinemáticas y dinámicas relevantes.",
            "Definir la cooperación entre AGV mediante grafos de comunicación local y condiciones de conectividad.",
            "Diseñar una estrategia nominal de control distribuido basada en consenso.",
            "Proponer una variante adaptativa o robusta frente a incertidumbre de masa, inercia o centro de masas de la carga.",
            "Implementar una plataforma de simulación reproducible en Python para comparar estrategias de control.",
            "Validar visualmente el comportamiento cooperativo en CoppeliaSim mediante una escena de bodega industrial.",
            "Evaluar las estrategias mediante métricas de seguimiento, formación, esfuerzo de control, convergencia y robustez cooperativa.",
            "Analizar los resultados para determinar en qué condiciones la cooperación adaptativa mejora o no a la estrategia nominal.",
        ],
    )

    add_p(doc, "2.3. Alcance del TFM", "Heading 2")
    add_p(
        doc,
        (
            "El alcance se define para mantener un equilibrio entre profundidad teórica y viabilidad técnica. "
            "El trabajo se concentrará en el control cooperativo distribuido para transporte de cargas en el "
            "plano. La plataforma experimental será simulada, pero la simulación se diseñará con criterios de "
            "reproducibilidad y trazabilidad de resultados."
        ),
    )
    add_p(doc, "El TFM incluye los siguientes elementos:")
    add_bullets(
        doc,
        [
            "modelo planar de un sistema compuesto por múltiples AGV y una carga heterogénea;",
            "grafos de comunicación local para representar cooperación entre agentes;",
            "controlador nominal distribuido basado en consenso;",
            "controlador adaptativo o robusto frente a incertidumbre inercial;",
            "escenarios de variación de masa, desplazamiento del centro de masas, degradación de comunicación y perturbaciones;",
            "simulación cuantitativa en Python y validación aplicada en CoppeliaSim.",
        ],
    )
    add_p(doc, "El TFM excluye explícitamente los siguientes elementos:")
    add_bullets(
        doc,
        [
            "construcción o ensayo con AGV físicos reales;",
            "percepción visual, SLAM o navegación autónoma completa;",
            "planificación global de rutas en bodega;",
            "identificación experimental completa de parámetros inerciales;",
            "certificación de seguridad industrial o integración con sistemas productivos reales.",
        ],
    )

    add_p(doc, "3. Hipótesis de partida", "Heading 1")
    add_p(
        doc,
        (
            "La hipótesis de partida orienta el diseño experimental y define una afirmación contrastable. En "
            "este TFM no se asume que la estrategia adaptativa será superior en todos los casos; se plantea que "
            "puede mejorar la consistencia del comportamiento cooperativo bajo determinadas formas de "
            "incertidumbre, con posibles costes en esfuerzo de control."
        ),
    )
    add_p(
        doc,
        (
            "Hipótesis principal: una estrategia de control cooperativo distribuido basada en consenso y con "
            "compensación adaptativa o robusta puede reducir la degradación del desempeño del transporte "
            "multi-AGV frente a incertidumbre inercial de la carga, en comparación con una estrategia nominal, "
            "siempre que el grafo de comunicación mantenga conectividad suficiente."
        ),
    )
    add_p(
        doc,
        (
            "Hipótesis secundarias: primero, la estrategia adaptativa reducirá el error de seguimiento y el error "
            "de formación en escenarios con variación de masa o desplazamiento del centro de masas. Segundo, "
            "la mejora de robustez puede implicar un incremento moderado del esfuerzo de control. Tercero, la "
            "ventaja del enfoque cooperativo disminuirá cuando la comunicación local se degrade de forma severa. "
            "Cuarto, una métrica de robustez cooperativa permitirá resumir la degradación relativa del desempeño "
            "entre controladores y escenarios."
        ),
    )

    add_p(doc, "4. Metodología", "Heading 1")
    add_p(
        doc,
        (
            "La metodología se diseña como una investigación aplicada, cuantitativa y experimental en simulación. "
            "El propósito es comparar estrategias de control bajo condiciones controladas, manteniendo la misma "
            "plataforma, las mismas trayectorias y las mismas perturbaciones para cada controlador. La metodología "
            "también incorpora una validación aplicada en CoppeliaSim, cuyo objetivo es visualizar y comprobar la "
            "coherencia robótica del comportamiento cooperativo."
        ),
    )
    add_p(
        doc,
        (
            "La lógica metodológica sigue la secuencia planteada para el TFM completo: definir el problema, "
            "analizar cómo se resuelve actualmente, identificar límites, proponer una plataforma, formular un "
            "método de control cooperativo, compararlo en Python, validarlo en CoppeliaSim y discutir los "
            "resultados. En la presente entrega se desarrolla hasta el diseño metodológico; los resultados y "
            "conclusiones se incorporarán en fases posteriores."
        ),
    )

    add_p(doc, "4.1. Diseño de investigación", "Heading 2")
    add_p(
        doc,
        (
            "El diseño de investigación es comparativo porque enfrenta una línea base nominal con una estrategia "
            "adaptativa o robusta. Es cuantitativo porque el desempeño se evaluará mediante métricas numéricas, "
            "y es experimental en simulación porque las variables independientes se modificarán de manera "
            "controlada. La comparación se realizará sobre escenarios equivalentes para evitar que las diferencias "
            "observadas provengan de condiciones iniciales, trayectorias o parámetros no controlados."
        ),
    )
    add_p(
        doc,
        (
            "La unidad de análisis será el sistema cooperativo completo, no un AGV individual. Esto implica que "
            "el desempeño se medirá sobre la carga, la formación y el conjunto de acciones de control. La decisión "
            "es coherente con la teoría central del trabajo: el objetivo no es optimizar un robot aislado, sino "
            "evaluar cómo la cooperación local permite resolver una tarea que supera la capacidad o conveniencia "
            "de un único agente."
        ),
    )
    add_p(
        doc,
        (
            "La investigación no tendrá como variable principal la calidad gráfica de la simulación, sino el "
            "comportamiento del sistema de control. Por ese motivo, las métricas se calculan sobre estados, "
            "errores y acciones de control. CoppeliaSim se incorporará como herramienta para comprobar la "
            "coherencia de la escena robótica y para comunicar de forma visual el transporte cooperativo, pero "
            "la comparación formal se basará en registros numéricos."
        ),
    )

    add_p(doc, "4.2. Fases del trabajo", "Heading 2")
    add_p(
        doc,
        (
            "El desarrollo se organizará en fases incrementales. Cada fase produce un artefacto verificable que "
            "alimenta la siguiente, evitando que la memoria dependa de resultados no reproducibles o de decisiones "
            "no documentadas."
        ),
    )
    add_numbered(
        doc,
        [
            "Revisión bibliográfica sobre consenso, control distribuido, cooperación multi-robot y transporte cooperativo.",
            "Definición del sistema robot-carga, supuestos de modelado y variables de estado.",
            "Diseño de grafos de comunicación local para representar la cooperación entre AGV.",
            "Implementación del controlador nominal basado en consenso.",
            "Diseño e implementación del controlador adaptativo o robusto.",
            "Configuración de experimentos reproducibles en Python.",
            "Construcción de una escena de validación en CoppeliaSim.",
            "Ejecución de escenarios, generación de métricas y producción de figuras.",
            "Análisis comparativo, discusión de limitaciones y redacción de conclusiones.",
        ],
    )

    add_p(doc, "4.3. Plataforma de simulación y validación", "Heading 2")
    add_p(
        doc,
        (
            "Python será la plataforma principal de simulación. En esta capa se implementarán el modelo, los "
            "controladores, las perturbaciones y las métricas. La elección responde a la necesidad de "
            "reproducibilidad: los experimentos podrán ejecutarse desde archivos de configuración, registrar "
            "resultados brutos y generar tablas o figuras sin intervención manual. CoppeliaSim se usará como "
            "plataforma de validación aplicada. Su función será representar el entorno de bodega, los AGV, la "
            "carga y el movimiento cooperativo, conectando el análisis de control con una escena robótica "
            "interpretable."
        ),
    )
    add_p(
        doc,
        (
            "La comunicación entre ambos niveles se planteará de forma pragmática. Python funcionará como banco "
            "de pruebas cuantitativo; CoppeliaSim reproducirá uno o varios escenarios representativos para "
            "mostrar el comportamiento cooperativo. Esta separación evita que las limitaciones gráficas o de "
            "configuración de CoppeliaSim condicionen la comparación científica, y al mismo tiempo evita que el "
            "TFM quede desconectado de una aplicación robótica visible."
        ),
    )
    add_p(
        doc,
        (
            "En el repositorio del TFM, cada experimento se documentará mediante archivos de configuración y "
            "salidas separadas en resultados brutos, resultados procesados, tablas y figuras. Esta organización "
            "permite repetir ejecuciones, revisar parámetros y asociar cada figura de la memoria con un "
            "experimento concreto. La trazabilidad es especialmente importante porque la evaluación final no "
            "debe depender de capturas aisladas, sino de una cadena reproducible desde configuración hasta "
            "resultado."
        ),
    )

    add_p(doc, "4.4. Método propuesto", "Heading 2")
    add_p(
        doc,
        (
            "El método nuevo se plantea como una arquitectura de control cooperativo distribuido. Cada AGV "
            "calcula su acción a partir de su estado local, el error respecto a la tarea cooperativa y la "
            "información recibida de sus vecinos en el grafo de comunicación. La línea base nominal utilizará "
            "un mecanismo de consenso para mantener coordinación y formación. La variante adaptativa o robusta "
            "incorporará un término de compensación para reducir el efecto de incertidumbres asociadas a la "
            "carga, especialmente masa efectiva, desplazamiento del centro de masas o perturbaciones externas."
        ),
    )
    add_p(
        doc,
        (
            "La contribución teórica-técnica no se presentará como una teoría universal, sino como una integración "
            "razonada de control, cooperación y evaluación. El aporte consiste en formular el problema con "
            "comunicación local, proponer una estrategia distribuida con compensación y medir su desempeño bajo "
            "escenarios que tensan la cooperación. De forma complementaria, se podrá definir un índice de robustez "
            "cooperativa para resumir la degradación relativa del sistema cuando cambia la carga o la conectividad."
        ),
    )
    add_p(
        doc,
        (
            "El diseño nominal servirá como referencia. Su función será coordinar los AGV mediante reglas de "
            "consenso y mantenimiento de formación bajo condiciones conocidas. El diseño adaptativo o robusto "
            "se introducirá para compensar desviaciones producidas por cargas diferentes a la nominal. La "
            "comparación entre ambos controladores permitirá distinguir entre desempeño nominal, sensibilidad "
            "a incertidumbre y coste de robustez."
        ),
    )

    add_p(doc, "4.5. Metodología de comparación", "Heading 2")
    add_p(
        doc,
        (
            "La comparación se realizará entre al menos dos controladores: una estrategia nominal distribuida y "
            "una estrategia adaptativa o robusta. Ambos se evaluarán con las mismas trayectorias, los mismos "
            "parámetros de simulación, las mismas condiciones iniciales y las mismas perturbaciones. Para cada "
            "escenario se registrarán trayectorias, errores, acciones de control y estado de la formación. Este "
            "criterio evita comparaciones anecdóticas y permite sostener conclusiones basadas en datos."
        ),
    )
    add_p(
        doc,
        (
            "En Python se ejecutará la batería completa de experimentos, mientras que en CoppeliaSim se validarán "
            "casos representativos: un escenario nominal y al menos un escenario con incertidumbre o perturbación. "
            "La comparación final combinará evidencia cuantitativa y evidencia visual aplicada, pero la decisión "
            "sobre desempeño se basará en las métricas calculadas de forma reproducible."
        ),
    )
    add_p(
        doc,
        (
            "La comparación se organizará por escenarios. En cada escenario se obtendrá una tabla resumen y un "
            "conjunto de figuras. La tabla permitirá revisar valores agregados, mientras que las figuras mostrarán "
            "trayectorias, evolución temporal de errores y acciones de control. En la discusión final se evitará "
            "afirmar superioridad absoluta si los datos muestran mejoras parciales o dependientes del escenario. "
            "Este criterio es relevante para mantener rigor académico y para que el TFM sea defendible ante el "
            "tribunal."
        ),
    )

    add_p(doc, "4.6. Variables, escenarios y métricas", "Heading 2")
    add_p(
        doc,
        (
            "Las variables se clasifican en independientes, dependientes y controladas. Esta clasificación permite "
            "mantener trazabilidad entre hipótesis, simulación y resultados esperados."
        ),
    )
    add_table(
        doc,
        "Tabla 2. Variables de estudio.",
        ["Tipo", "Variables"],
        [
            [
                "Independientes",
                "Tipo de controlador, masa de la carga, centro de masas, perturbaciones externas y topología del grafo.",
            ],
            [
                "Dependientes",
                "Error de seguimiento, error de formación, esfuerzo de control, tiempo de convergencia y robustez cooperativa.",
            ],
            [
                "Controladas",
                "Trayectoria de referencia, número de AGV, condiciones iniciales, tiempo de simulación, paso temporal y semilla.",
            ],
        ],
    )
    add_table(
        doc,
        "Tabla 3. Escenarios de evaluación.",
        ["Escenario", "Propósito"],
        [
            ["Caso nominal", "Verificar el comportamiento base de la arquitectura sin incertidumbre severa."],
            ["Variación de masa", "Evaluar sensibilidad frente a cargas más ligeras o pesadas que la nominal."],
            ["Desplazamiento del centro de masas", "Analizar el efecto de una distribución de peso no centrada."],
            ["Degradación de comunicación", "Medir impacto de pérdida parcial de enlaces o reducción de conectividad."],
            ["Perturbaciones externas", "Evaluar rechazo de disturbios acotados durante el transporte."],
        ],
    )
    add_table(
        doc,
        "Tabla 4. Métricas de comparación.",
        ["Métrica", "Interpretación"],
        [
            ["Error de seguimiento", "Diferencia entre trayectoria deseada y trayectoria observada de la carga."],
            ["Error de formación", "Desviación respecto a la geometría cooperativa esperada entre AGV."],
            ["Esfuerzo de control", "Magnitud acumulada de las acciones aplicadas por los robots."],
            ["Tiempo de convergencia", "Tiempo requerido para entrar y permanecer dentro de una banda de error."],
            ["Robustez cooperativa", "Degradación relativa del desempeño ante incertidumbre o perturbación."],
        ],
    )

    add_p(doc, "4.7. Limitaciones metodológicas", "Heading 2")
    add_p(
        doc,
        (
            "El alcance del TFM queda limitado a simulación y validación virtual. No se realizarán pruebas con AGV "
            "físicos, ni se incorporará percepción real, SLAM, visión artificial o planificación global de rutas. "
            "Tampoco se pretende resolver la identificación completa de los parámetros inerciales de la carga. "
            "Estas exclusiones son deliberadas: permiten concentrar el trabajo en control cooperativo, consenso, "
            "comunicación local y evaluación de robustez."
        ),
    )
    add_p(
        doc,
        (
            "Otra limitación prevista es que CoppeliaSim se usará como entorno de validación aplicada y no como "
            "única fuente de resultados. La simulación física detallada puede introducir efectos que no forman "
            "parte del modelo matemático simplificado, por lo que la interpretación deberá distinguir entre "
            "validación conceptual, resultados cuantitativos y visualización robótica. Esta separación será "
            "necesaria para mantener coherencia entre teoría, simulación y conclusiones."
        ),
    )

    add_p(doc, "4.8. Criterios de aceptación y trazabilidad", "Heading 2")
    add_p(
        doc,
        (
            "El método se considerará satisfactorio si permite ejecutar todos los escenarios definidos, generar "
            "automáticamente métricas comparables y producir evidencia suficiente para aceptar, matizar o rechazar "
            "la hipótesis de partida. No se exigirá que el controlador adaptativo mejore todos los indicadores en "
            "todos los casos; sí se exigirá que las diferencias observadas puedan explicarse técnicamente."
        ),
    )
    add_p(
        doc,
        (
            "La trazabilidad del trabajo se garantizará mediante tres mecanismos. Primero, cada escenario tendrá "
            "un archivo de configuración con parámetros explícitos. Segundo, los resultados se almacenarán de "
            "forma separada para evitar reescrituras manuales. Tercero, las figuras y tablas se generarán desde "
            "los resultados procesados. Esta disciplina permitirá que la memoria final no sea solo descriptiva, "
            "sino verificable."
        ),
    )
    add_p(
        doc,
        (
            "Para la entrega de la Tarea 1, el resultado esperado es dejar completamente planteado el problema, "
            "la pregunta de investigación, los objetivos, las hipótesis, el alcance y la metodología de comparación. "
            "En las fases siguientes se completarán el marco teórico ampliado, la implementación, los experimentos, "
            "los resultados y las conclusiones."
        ),
    )

    add_p(doc, "Referencias bibliográficas", "Heading 1")
    refs = [
        "Bullo, F., Cortés, J., & Martínez, S. (2009). Distributed control of robotic networks: A mathematical approach to motion coordination algorithms. Princeton University Press.",
        "Cao, Y. U., Fukunaga, A. S., & Kahng, A. B. (1997). Cooperative mobile robotics: Antecedents and directions. Autonomous Robots, 4, 7-27. https://doi.org/10.1023/A:1008855018923",
        "Coppelia Robotics. (s. f.). CoppeliaSim user manual. https://manual.coppeliarobotics.com/",
        "Dudek, G., Jenkin, M. R. M., Milios, E., & Wilkes, D. (1996). A taxonomy for multi-agent robotics. Autonomous Robots, 3, 375-397. https://doi.org/10.1007/BF00240651",
        "Olfati-Saber, R., Fax, J. A., & Murray, R. M. (2007). Consensus and cooperation in networked multi-agent systems. Proceedings of the IEEE, 95(1), 215-233. https://doi.org/10.1109/JPROC.2006.887293",
        "Parker, L. E. (2008). Multiple mobile robot systems. En B. Siciliano & O. Khatib (Eds.), Springer handbook of robotics (pp. 921-941). Springer.",
        "Ren, W., & Beard, R. W. (2008). Distributed consensus in multi-vehicle cooperative control: Theory and applications. Springer. https://doi.org/10.1007/978-1-84800-015-5",
        "Yufka, A., & Ozkan, M. (2015). Formation-based control scheme for cooperative transportation by multiple mobile robots. International Journal of Advanced Robotic Systems, 12(9). https://doi.org/10.5772/60972",
    ]
    for ref in refs:
        add_p(doc, ref)

    # Ensure basic page layout and body style.
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.start_type = WD_SECTION_START.NEW_PAGE
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)

    doc.save(OUT_DOCX)
    strip_template_toc_field(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
