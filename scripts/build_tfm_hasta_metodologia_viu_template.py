"""Build the TFM draft through methodology using the official VIU template."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import subprocess
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\walla\Downloads\Plantilla memoria TFM_ MROB.docx")
SOURCE_DOCX = ROOT / "docs" / "doc-01-proposal" / "TFM_MayorgaTaborda_hasta_metodologia.docx"
OUT_DOCX = ROOT / "docs" / "doc-01-proposal" / "TFM_MayorgaTaborda_hasta_metodologia_VIU.docx"
LATEX_SOURCE = ROOT / "docs" / "doc-01-proposal" / "main_hasta_metodologia.tex"
BIB_SOURCE = ROOT / "docs" / "doc-01-proposal" / "references.bib"

TITLE = (
    "Control distribuido basado en dinámicas poblacionales para la formación "
    "adaptativa de coaliciones en transporte cooperativo multi-AGV"
)
AUTHOR = "Jorge Luis Mayorga Taborda"
TUTOR = "José Ignacio Iñíguez Amigot"
DATE_TEXT = "27 de mayo de 2026"
EDITION = "2025-26"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_font(paragraph, size: int = 12) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)


def setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def normalize_text(text: str) -> str:
    replacements = {
        "actuen": "actúen",
        "metodos": "métodos",
        "posicion": "posición",
        "estan": "están",
        "posicion": "posición",
        "formara": "formará",
        "degradacion": "degradación",
        "afectara": "afectará",
        "reducira": "reducirá",
        "integracion": "integración",
        "ofrecera": "ofrecerá",
        "contrastaran": "contrastarán",
        "metricas": "métricas",
        "funcion": "función",
        "cohesion": "cohesión",
        "analisis": "análisis",
        "campana": "campaña",
        "numerica": "numérica",
        "integracion": "integración",
        "fisico": "físico",
        "delimitacion": "delimitación",
        "seguira": "seguirá",
        "experimentacion": "experimentación",
        "debera": "deberá",
        "cerrara": "cerrará",
        "formulacion": "formulación",
        "tendran": "tendrán",
        "distribucion": "distribución",
        "representacion": "representación",
        "implementaran": "implementarán",
        "todavia": "todavía",
        "asignara": "asignará",
        "aumentara": "aumentará",
        "disminuira": "disminuirá",
        "este cerca": "esté cerca",
        "esta expresion": "esta expresión",
        "extinguidas": "extinguidas",
        "servira": "servirá",
        "Servira": "Servirá",
        "resolvera": "resolverá",
        "saturada": "saturada",
        "este saturada": "esté saturada",
        "expresion": "expresión",
        "participa": "participa",
        "atraccion": "atracción",
        "repulsion": "repulsión",
        "realizara": "realizará",
        "configuracion": "configuración",
        "preferencias": "preferencias",
        "actualizaran": "actualizarán",
        "version": "versión",
        "proyeccion": "proyección",
        "sustituira": "sustituirá",
        "servira": "servirá",
        "que carga": "qué carga",
        "que peso": "qué peso",
        "visualizacion": "visualización",
        "separacion": "separación",
        "organizara": "organizará",
        "limitara": "limitará",
        "reportaran": "reportarán",
        "parametrica": "paramétrica",
        "comparara": "comparará",
        "definira": "definirá",
        "adaptacion": "adaptación",
        "evitara": "evitará",
        "encontro": "encontró",
        "Proposito": "Propósito",
        "Resultado esperado": "Resultado esperado",
        "Entregadas": "Entregadas",
        "numero": "número",
        "graficas": "gráficas",
        "ejecutara": "ejecutará",
        "mostraran": "mostrarán",
        "descentralizacion": "descentralización",
        "sensibilidad parametrica": "sensibilidad paramétrica",
        "dificil": "difícil",
        "basica": "básica",
        "homogeneo": "homogéneo",
        "Homogeneo": "Homogéneo",
        "Coalicion": "Coalición",
        "Metrica": "Métrica",
        "Definicion": "Definición",
        "Proposito": "Propósito",
        "metrica": "métrica",
        "minima": "mínima",
        "maxima": "máxima",
        "cinetico": "cinético",
        "separara": "separará",
        "friccion": "fricción",
        "pendiente sobre": "reservado para trabajo futuro sobre",
    }
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(source, target)
    cleanup = {
        "configuraciónes": "configuraciones",
        "posiciónes": "posiciones",
        "En está expresión": "En esta expresión",
        "está expresión": "esta expresión",
        "está saturada": "esté saturada",
        "propuestá": "propuesta",
        "sistemás": "sistemas",
    }
    for source, target in cleanup.items():
        text = text.replace(source, target)
    return text


def normalize_document_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        updated = normalize_text(paragraph.text)
        if updated != paragraph.text:
            paragraph.text = updated
            set_font(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    updated = normalize_text(paragraph.text)
                    if updated != paragraph.text:
                        paragraph.text = updated
                        set_font(paragraph, 10)


def add_p(doc: Document, text: str, style: str = "Normal", align=None):
    paragraph = doc.add_paragraph(normalize_text(text), style=style)
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    style = "Heading 1" if level == 1 else "Heading 2"
    paragraph = doc.add_paragraph(normalize_text(text), style=style)
    set_font(paragraph)
    return paragraph


def add_list_item(doc: Document, text: str, numbered: bool, index: int | None = None) -> None:
    prefix = f"{index}. " if numbered and index is not None else "• "
    paragraph = doc.add_paragraph(prefix + normalize_text(text), style="List Paragraph")
    set_font(paragraph)


def append_table_from_source(target_doc: Document, source_table) -> None:
    target_doc._body._element.append(deepcopy(source_table._tbl))
    target_doc.add_paragraph("")


def strip_template_dynamic_blocks(docx_path: Path) -> None:
    """Remove original TOC/index fields from the VIU template body."""

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path, "r") as src:
        files = {name: src.read(name) for name in src.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    removed = 0
    for sdt in list(root.xpath(".//w:sdt", namespaces=ns)):
        text = "".join(sdt.xpath(".//w:t/text() | .//w:instrText/text()", namespaces=ns))
        if "TOC" in text or "Índice" in text or "Figura 1" in text or "Tabla 1" in text:
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)
                removed += 1
    if removed:
        files["word/document.xml"] = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone="yes",
        )
        tmp_path = docx_path.with_suffix(".tmp.docx")
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as dst:
            for name, data in files.items():
                dst.writestr(name, data)
        tmp_path.replace(docx_path)


def fill_cover(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "TÍTULO:":
            paragraph.text = f"TÍTULO:\n{TITLE}"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph)
            for run in paragraph.runs:
                run.bold = True
        elif text == "Alumno/a:":
            paragraph.text = f"Alumno/a: {AUTHOR}"
            set_font(paragraph)
        elif text == "Director/a:":
            paragraph.text = f"Director/a: {TUTOR}"
            set_font(paragraph)
        elif text.startswith("Edición:"):
            paragraph.text = f"Edición: {EDITION}\nFecha: {DATE_TEXT}"
            set_font(paragraph)


def clear_template_body(doc: Document) -> None:
    start = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Resumen":
            start = index
            break
    if start is None:
        return
    for paragraph in list(doc.paragraphs[start:]):
        remove_paragraph(paragraph)


def source_parts():
    source = Document(SOURCE_DOCX)
    paragraphs = [p.text.strip() for p in source.paragraphs]
    paragraphs = [text for text in paragraphs if text]

    resumen_start = paragraphs.index("Resumen")
    contenido_start = paragraphs.index("Contenido")
    intro_start = paragraphs.index("1. Introducción", contenido_start + 1)
    refs_start = paragraphs.index("Referencias bibliográficas", intro_start + 1)

    resumen = paragraphs[resumen_start + 1 : contenido_start]
    body = paragraphs[intro_start:refs_start]
    references = paragraphs[refs_start + 1 :]
    return resumen, body, references, source.tables


def build_doc() -> None:
    if LATEX_SOURCE.exists():
        subprocess.run(
            [
                "pandoc",
                str(LATEX_SOURCE),
                "--from",
                "latex",
                "--to",
                "docx",
                "--number-sections",
                "--toc",
                "--metadata",
                f"title={TITLE}",
                "--metadata",
                f"author={AUTHOR}",
                "--metadata",
                f"date={DATE_TEXT}",
                "--metadata",
                "reference-section-title=Referencias bibliográficas",
                "--citeproc",
                "--bibliography",
                str(BIB_SOURCE),
                "--output",
                str(OUT_DOCX),
            ],
            cwd=ROOT,
            check=True,
        )
        doc = Document(OUT_DOCX)
        setup_page(doc)
        doc.save(OUT_DOCX)
        print(OUT_DOCX)
        return

    doc = Document(TEMPLATE)
    fill_cover(doc)
    clear_template_body(doc)

    resumen, body, references, tables = source_parts()

    doc.add_page_break()
    add_p(doc, "Resumen", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    for paragraph in resumen:
        add_p(doc, paragraph)

    doc.add_page_break()
    add_p(doc, "Contenido", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    for item in [
        "1. Introducción",
        "2. Planteamiento del problema y justificación",
        "3. Pregunta de investigación e hipótesis",
        "4. Objetivos",
        "5. Alcance y limitaciones",
        "6. Metodología",
        "Referencias bibliográficas",
    ]:
        add_p(doc, item)

    table_index = 0
    numbered_until: str | None = None
    bullet_until: str | None = None
    number_index = 1
    for text in body:
        if text.startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ")):
            add_heading(doc, text, 1)
            numbered_until = None
            bullet_until = None
            number_index = 1
        elif text.startswith(("4.1.", "4.2.", "5.1.", "5.2.", "6.1.", "6.2.", "6.3.", "6.4.", "6.5.", "6.6.", "6.7.")):
            add_heading(doc, text, 2)
            numbered_until = "5." if text.startswith("4.2.") else None
            bullet_until = None
            number_index = 1
            if text.startswith("6.6.") and table_index < len(tables):
                append_table_from_source(doc, tables[table_index])
                table_index += 1
        elif text == "Hipótesis de trabajo:":
            add_p(doc, text)
            numbered_until = "Las hipótesis"
            number_index = 1
        elif text.startswith("Las hipótesis"):
            numbered_until = None
            add_p(doc, text)
        elif text.startswith("El alcance se define"):
            add_p(doc, text)
            bullet_until = "5.2."
        elif text.startswith("La campana experimental") or text.startswith("La campaña experimental"):
            add_p(doc, text)
            if table_index < len(tables):
                append_table_from_source(doc, tables[table_index])
                table_index += 1
        elif numbered_until and not text.startswith(numbered_until):
            add_list_item(doc, text, numbered=True, index=number_index)
            number_index += 1
        elif bullet_until:
            add_list_item(doc, text, numbered=False)
        else:
            add_p(doc, text)

    add_p(doc, "Referencias bibliográficas", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    for reference in references:
        paragraph = add_p(doc, reference)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)

    normalize_document_text(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    strip_template_dynamic_blocks(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
